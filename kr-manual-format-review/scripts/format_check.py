#!/usr/bin/env python3
"""한국어 매뉴얼 서식 검수 스크립트.

검수(기본) / 기계적 수정(--fix) 두 모드를 지원한다.
구조적 수정(제목 레벨, 목록 numId 통합, 콜아웃 표 변환)은 SKILL.md 절차에 따라
에이전트가 직접 수행한다.

사용:
    python3 format_check.py --file 매뉴얼.docx                  # 전체 검수
    python3 format_check.py --file 매뉴얼.docx --section "HA 관리"
    python3 format_check.py --file 매뉴얼.docx --fix            # 기계적 수정
"""

import argparse
import re
import sys
import zipfile
from xml.etree import ElementTree as ET

import docx

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
STANDARD_FONT = "Pretendard"
BAD_FONT_HINT = "JP"  # 'Pretendard JP Variable' 등

HEADING_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)[\.\s]")
SPACING_TYPOS = ["선택후", "] 에서", "] 를"]


def section_range(paras, section):
    """H1 제목명으로 검수 범위(파라그래프 인덱스)를 구한다."""
    if not section:
        return 0, len(paras) - 1
    start = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 1" and p.text.strip().replace("\xa0", " ") == section:
            start = i
            break
    if start is None:
        sys.exit(f"H1 '{section}' 섹션을 찾을 수 없습니다.")
    end = len(paras) - 1
    for i in range(start + 1, len(paras)):
        if paras[i].style.name == "Heading 1":
            end = i - 1
            break
    return start, end


def load_numbering(path):
    """numId -> (numFmt, lvlText) of lvl0."""
    with zipfile.ZipFile(path) as z:
        if "word/numbering.xml" not in z.namelist():
            return {}
        root = ET.fromstring(z.read("word/numbering.xml"))
    abstracts = {}
    for a in root.findall(f"{W}abstractNum"):
        aid = a.get(f"{W}abstractNumId")
        for lvl in a.findall(f"{W}lvl"):
            if lvl.get(f"{W}ilvl") == "0":
                fmt = lvl.find(f"{W}numFmt")
                txt = lvl.find(f"{W}lvlText")
                abstracts[aid] = (
                    fmt.get(f"{W}val") if fmt is not None else "?",
                    txt.get(f"{W}val") if txt is not None else "?",
                )
    out = {}
    for n in root.findall(f"{W}num"):
        nid = n.get(f"{W}numId")
        aid = n.find(f"{W}abstractNumId").get(f"{W}val")
        out[nid] = abstracts.get(aid, ("?", "?"))
    return out


def para_numid(p):
    numpr = p._element.find(f"{W}pPr/{W}numPr")
    if numpr is None:
        return None
    nid = numpr.find(f"{W}numId")
    return nid.get(f"{W}val") if nid is not None else None


def heading_level(style_name):
    m = re.match(r"Heading (\d)", style_name)
    return int(m.group(1)) if m else None


def check(doc, path, start, end):
    paras = doc.paragraphs
    num_defs = load_numbering(path)
    issues = []

    def add(idx, kind, msg):
        issues.append((idx, kind, msg))

    prev_numids = []
    for i in range(start, end + 1):
        p = paras[i]
        text = p.text
        stripped = text.strip()
        st = p.style.name
        lvl = heading_level(st)

        if lvl:
            # 제목에 자동 번호 매기기 금지
            if para_numid(p):
                add(i, "heading-autonum", f"제목에 자동 번호(numPr): [{st}] {stripped[:50]}")
            # 번호 깊이 vs Heading 레벨
            m = HEADING_NUM_RE.match(stripped)
            if m and lvl >= 2:
                depth = m.group(1).count(".") + 1
                expected = depth + 1  # '1.'=H2, '1.1'=H3 ...
                if expected != lvl and expected <= 4:
                    add(i, "heading-level",
                        f"번호 깊이({m.group(1)})와 레벨 불일치: [{st}] {stripped[:50]}")
            elif not m and lvl >= 3:
                add(i, "heading-nonum", f"제목 번호 누락 추정: [{st}] {stripped[:50]}")
            if "삭제/수정" in stripped:
                add(i, "wording", f"'삭제/수정' → '수정/삭제': {stripped[:50]}")
            if re.search(r"상세 -\s", stripped):
                add(i, "wording", f"'상세 -' → '상세 정보 -': {stripped[:50]}")

        # 공통: nbsp / 비표준 폰트 / 오타
        if "\xa0" in text:
            add(i, "nbsp", f"특수 공백 {text.count(chr(160))}개: {stripped[:40]}")
        for r in p.runs:
            rPr = r.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                if any(BAD_FONT_HINT in (v or "") for v in rPr.rFonts.attrib.values()):
                    add(i, "font", f"비표준 폰트 run: {stripped[:40]}")
                    break
        for typo in SPACING_TYPOS:
            if typo in text:
                add(i, "spacing", f"'{typo}' 발견: {stripped[:50]}")
        if not lvl and "삭제/수정" in text:
            add(i, "wording", f"'삭제/수정' → '수정/삭제': {stripped[:60]}")
        if stripped.endswith(("습니다", "합니다", "됩니다")):
            add(i, "period", f"마침표 누락: …{stripped[-25:]}")

        # 본문 목록 번호 형식
        nid = para_numid(p)
        if nid and not lvl:
            fmt, lvltext = num_defs.get(nid, ("?", "?"))
            if lvltext.endswith("."):
                add(i, "list-format", f"목록 'N.' 형식(관례는 'N)'): numId={nid} {stripped[:40]}")
            prev_numids.append((i, nid))

        # 콜아웃이 일반 문단으로 들어간 경우
        if stripped in ("주의 사항", "참고 사항") and not lvl:
            add(i, "callout", f"'{stripped}'이 일반 문단 — 콜아웃 표로 변환 필요")

    # 인접 목록 항목이 제각각 numId를 쓰는 패턴 (start override 흉내)
    for k in range(len(prev_numids) - 1):
        (i1, n1), (i2, n2) = prev_numids[k], prev_numids[k + 1]
        if n1 != n2 and i2 - i1 <= 3:  # 이미지 1~2장 끼어도 같은 절차로 간주
            add(i2, "list-split", f"인접 목록이 다른 numId({n1}→{n2}) — 단일 리스트로 통합 검토")

    return issues


def mechanical_fix(doc, start, end):
    """nbsp→공백, JP 폰트→Pretendard, 이중 공백 정리. 수정 건수 반환."""
    paras = doc.paragraphs
    counts = {"nbsp": 0, "font": 0, "dspace": 0}
    qn = docx.oxml.ns.qn
    for i in range(start, end + 1):
        p = paras[i]
        is_heading = p.style.name.startswith("Heading")
        for r in p.runs:
            if "\xa0" in r.text:
                counts["nbsp"] += r.text.count("\xa0")
                r.text = r.text.replace("\xa0", " ")
            if "  " in r.text:
                counts["dspace"] += 1
                r.text = re.sub(r" {2,}", " ", r.text)
            rPr = r.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                rf = rPr.rFonts
                if any(BAD_FONT_HINT in (v or "") for v in rf.attrib.values()):
                    counts["font"] += 1
                    if is_heading:
                        rPr.remove(rf)  # 스타일 상속
                    else:
                        for attr in list(rf.attrib):
                            if BAD_FONT_HINT in rf.attrib[attr]:
                                if attr == qn("w:eastAsia"):
                                    del rf.attrib[attr]
                                else:
                                    rf.attrib[attr] = STANDARD_FONT
    return counts


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--file", required=True)
    ap.add_argument("--section", help="H1 제목명으로 범위 한정 (예: 'HA 관리')")
    ap.add_argument("--fix", action="store_true", help="기계적 수정 적용 후 저장")
    args = ap.parse_args()

    doc = docx.Document(args.file)
    start, end = section_range(doc.paragraphs, args.section)
    scope = args.section or "문서 전체"
    print(f"# 검수 범위: {scope} (paragraph {start}–{end})\n")

    if args.fix:
        counts = mechanical_fix(doc, start, end)
        doc.save(args.file)
        print(f"기계적 수정 완료: nbsp {counts['nbsp']}개, "
              f"비표준 폰트 run {counts['font']}개, 이중 공백 {counts['dspace']}건")
        doc = docx.Document(args.file)  # 재검수

    issues = check(doc, args.file, start, end)
    if not issues:
        print("발견된 서식 이슈 없음.")
        return
    by_kind = {}
    for idx, kind, msg in issues:
        by_kind.setdefault(kind, []).append((idx, msg))
    for kind in sorted(by_kind):
        print(f"\n## {kind} ({len(by_kind[kind])}건)")
        for idx, msg in by_kind[kind]:
            print(f"  p{idx:5d}  {msg}")
    print(f"\n총 {len(issues)}건 — 구조적 항목은 SKILL.md MODE B 절차로 수정하세요.")


if __name__ == "__main__":
    main()
